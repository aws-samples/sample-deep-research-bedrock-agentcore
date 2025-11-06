"""Word Document Merger Utilities

Merge multiple Word documents into a single final report without using LLM.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List
import re


def append_markdown_as_document(master: Document, markdown_path: str):
    """
    Convert markdown file to Word document paragraphs and append to master.

    Args:
        master: Master document to append to
        markdown_path: Path to markdown file
    """
    with open(markdown_path, 'r', encoding='utf-8') as f:
        markdown_content = f.read()

    lines = markdown_content.split('\n')

    for line in lines:
        line = line.rstrip()

        if not line:
            continue

        # Heading level 1
        if line.startswith('# '):
            master.add_heading(line[2:], level=1)

        # Heading level 2
        elif line.startswith('## '):
            master.add_heading(line[3:], level=2)

        # Heading level 3
        elif line.startswith('### '):
            master.add_heading(line[4:], level=3)

        # Bullet list
        elif line.startswith('- ') or line.startswith('* '):
            master.add_paragraph(line[2:], style='List Bullet')

        # Numbered list
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            master.add_paragraph(text, style='List Number')

        # Regular paragraph
        else:
            master.add_paragraph(line)


def append_document(master_doc: Document, source_doc_path: str) -> None:
    """
    Append all content from source document to master document.

    This preserves all formatting, styles, tables, images, etc.

    Args:
        master_doc: Master Document object to append to
        source_doc_path: Path to source document file
    """
    source_doc = Document(source_doc_path)

    # Copy all body elements (paragraphs, tables, sections, etc.)
    for element in source_doc.element.body:
        master_doc.element.body.append(element)


def merge_dimension_documents(
    dimension_doc_paths: List[str],
    output_path: str,
    title: str,
    executive_summary: str = None,
    introduction: str = None,
    conclusion: str = None,
) -> str:
    """
    Merge multiple dimension documents into a single final report.

    Structure:
    1. Title
    2. Executive Summary (if provided)
    3. Introduction (if provided)
    4. Dimension sections (from separate documents)
    5. Conclusion (if provided)
    6. References (collected from all documents)

    Args:
        dimension_doc_paths: List of paths to dimension document files
        output_path: Path for output merged document
        title: Report title
        executive_summary: Optional executive summary text
        introduction: Optional introduction text
        conclusion: Optional conclusion text

    Returns:
        Path to merged document
    """
    # Create master document
    master = Document()

    # Add title
    title_para = master.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add executive summary if provided
    if executive_summary:
        master.add_page_break()
        master.add_heading("Executive Summary", level=1)
        master.add_paragraph(executive_summary)

    # Add introduction if provided
    if introduction:
        master.add_page_break()
        master.add_heading("Introduction", level=1)
        master.add_paragraph(introduction)

    # Append each dimension document
    for doc_path in dimension_doc_paths:
        master.add_page_break()

        # Check if it's markdown or docx
        if doc_path and doc_path.endswith('.md'):
            # Convert markdown to Word paragraphs
            append_markdown_as_document(master, doc_path)
        elif doc_path and doc_path.endswith('.docx'):
            # Append existing Word document
            append_document(master, doc_path)
        else:
            print(f"   ⚠ Skipping unknown file type: {doc_path}")

    # Add conclusion if provided
    if conclusion:
        master.add_page_break()
        master.add_heading("Conclusion", level=1)
        master.add_paragraph(conclusion)

    # Save merged document
    master.save(output_path)

    return output_path


def collect_references_from_documents(doc_paths: List[str]) -> List[str]:
    """
    Collect all references from multiple documents.

    Looks for paragraphs in "References" sections and extracts citations.

    Args:
        doc_paths: List of document paths

    Returns:
        List of unique reference strings
    """
    references = set()

    for doc_path in doc_paths:
        # Skip markdown files - references will be included when converted
        if doc_path.endswith('.md'):
            continue

        doc = Document(doc_path)
        in_references = False

        for para in doc.paragraphs:
            # Check if we've entered references section
            if para.style.name.startswith('Heading') and 'reference' in para.text.lower():
                in_references = True
                continue

            # Collect references
            if in_references and para.text.strip():
                # Stop at next major heading
                if para.style.name.startswith('Heading 1'):
                    break

                references.add(para.text.strip())

    return sorted(list(references))


def add_references_section(doc: Document, references: List[str]) -> None:
    """
    Add references section to document.

    Args:
        doc: Document object
        references: List of reference strings
    """
    doc.add_page_break()
    doc.add_heading("References", level=1)

    for i, ref in enumerate(references, 1):
        para = doc.add_paragraph(f"{i}. {ref}")
        para.style = 'List Number'
