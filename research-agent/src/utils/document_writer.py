"""Word document writer utilities using python-docx

This module provides utilities for creating structured Word documents
from research results with proper formatting and citations.
"""

import re
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def create_research_document(title: str) -> Document:
    """
    Create a new Word document with predefined styles for research reports.

    Args:
        title: Document title

    Returns:
        Configured Document object
    """
    doc = Document()

    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def add_section_heading(doc: Document, text: str, level: int = 1):
    """
    Add a heading to the document.

    Args:
        doc: Document object
        text: Heading text
        level: Heading level (1-3)
    """
    heading = doc.add_heading(text, level=level)
    return heading


def parse_markdown_to_word(doc: Document, markdown_text: str):
    """
    Parse markdown text and add to Word document with proper formatting.

    Supports:
    - Headings (##, ###, ####)
    - Bold (**text** or __text__)
    - Italic (*text* or _text_)
    - Citations [Author et al., Year]
    - Lists (-, *, 1.)
    - Code blocks (```)

    Args:
        doc: Document object
        markdown_text: Markdown formatted text
    """
    lines = markdown_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block detection
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block - add as monospace paragraph
                code_text = '\n'.join(code_lines)
                code_para = doc.add_paragraph(code_text)
                code_para.style = 'Intense Quote'
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Headings
        if line.startswith('####'):
            add_section_heading(doc, line[4:].strip(), level=3)
        elif line.startswith('###'):
            add_section_heading(doc, line[3:].strip(), level=2)
        elif line.startswith('##'):
            add_section_heading(doc, line[2:].strip(), level=1)

        # Bullet lists
        elif line.strip().startswith(('- ', '* ')):
            text = line.strip()[2:]
            para = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(para, text)

        # Numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            para = doc.add_paragraph(style='List Number')
            _add_formatted_text(para, text)

        # Regular paragraph
        else:
            para = doc.add_paragraph()
            _add_formatted_text(para, line)

        i += 1


def _add_formatted_text(paragraph, text: str):
    """
    Add text to paragraph with inline formatting (bold, italic, citations).

    Args:
        paragraph: Paragraph object
        text: Text with markdown formatting
    """
    # Parse inline formatting
    parts = []
    current_pos = 0

    # Pattern for bold, italic, and citations
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|_[^_]+_|\[[^\]]+\])'

    for match in re.finditer(pattern, text):
        # Add text before match
        if match.start() > current_pos:
            parts.append(('normal', text[current_pos:match.start()]))

        matched_text = match.group(0)

        # Bold
        if matched_text.startswith('**') or matched_text.startswith('__'):
            inner_text = matched_text[2:-2]
            parts.append(('bold', inner_text))

        # Italic
        elif matched_text.startswith('*') or matched_text.startswith('_'):
            inner_text = matched_text[1:-1]
            parts.append(('italic', inner_text))

        # Citation
        elif matched_text.startswith('['):
            parts.append(('citation', matched_text))

        current_pos = match.end()

    # Add remaining text
    if current_pos < len(text):
        parts.append(('normal', text[current_pos:]))

    # Add runs to paragraph
    for format_type, content in parts:
        run = paragraph.add_run(content)

        if format_type == 'bold':
            run.bold = True
        elif format_type == 'italic':
            run.italic = True
        elif format_type == 'citation':
            run.italic = True
            run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # Blue color for citations


def add_executive_summary(doc: Document, summary: str):
    """
    Add executive summary section.

    Args:
        doc: Document object
        summary: Summary text (can include markdown)
    """
    add_section_heading(doc, "Executive Summary", level=1)

    # Add summary box with light gray background
    summary_para = doc.add_paragraph()
    summary_para.style = 'Intense Quote'
    parse_markdown_to_word(doc, summary)


def add_dimension_section(doc: Document, dimension: str, content: str):
    """
    Add a dimension section with content.

    Args:
        doc: Document object
        dimension: Dimension name
        content: Section content (markdown formatted)
    """
    add_section_heading(doc, dimension, level=1)
    parse_markdown_to_word(doc, content)
    doc.add_page_break()


def add_references(doc: Document, citations: List[str]):
    """
    Add references section.

    Args:
        doc: Document object
        citations: List of citation strings
    """
    add_section_heading(doc, "References", level=1)

    for i, citation in enumerate(citations, 1):
        para = doc.add_paragraph(f"{i}. {citation}")
        para.style = 'List Number'


def save_document(doc: Document, filename: str):
    """
    Save document to file.

    Args:
        doc: Document object
        filename: Output filename
    """
    doc.save(filename)
    return filename


def extract_citations_from_markdown(text: str) -> List[str]:
    """
    Extract all citations from markdown text.

    Citations are in format [Author et al., Year, Source]

    Args:
        text: Markdown text

    Returns:
        List of unique citations
    """
    pattern = r'\[([^\]]+(?:et al\.|[A-Z][a-z]+)[^\]]*(?:19|20)\d{2}[^\]]*)\]'
    matches = re.findall(pattern, text)

    # Return unique citations preserving order
    seen = set()
    unique_citations = []
    for match in matches:
        if match not in seen:
            seen.add(match)
            unique_citations.append(match)

    return unique_citations
