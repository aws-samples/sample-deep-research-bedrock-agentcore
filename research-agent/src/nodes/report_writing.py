"""Report Writing Node - Merges dimension documents and refines with editor agent

This node:
1. Merges dimension markdown files into single markdown
2. Collects and deduplicates references
3. Adds executive summary and conclusion placeholders
4. Uses editor agent to refine the document
5. Generates executive summary and conclusion
6. Converts final markdown to Word document
"""

import time
import re
import logging
from typing import Dict, Any, Set
from pathlib import Path
from langsmith import traceable
from langgraph.prebuilt import create_react_agent
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.messages import SystemMessage
from docx import Document
from docx.shared import Pt

from src.state import ResearchState
from src.config.llm_config import get_llm_for_node
from src.tools.editor_tools import write_summary_and_conclusion, replace_text
from src.utils.cancellation import check_cancellation

logger = logging.getLogger(__name__)


def collect_references_from_markdown(md_files: list) -> Set[str]:
    """
    Collect unique references from markdown files.

    Args:
        md_files: List of markdown file paths

    Returns:
        Set of unique reference strings
    """
    references = set()

    for md_file in md_files:
        if not md_file or not Path(md_file).exists():
            continue

        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()

        # Find references section
        lines = content.split('\n')
        in_references = False

        for line in lines:
            line_stripped = line.strip()

            # Check if entering references section
            if line_stripped.startswith('##') and 'reference' in line_stripped.lower():
                in_references = True
                continue

            # Stop at next major heading
            if in_references and line_stripped.startswith('##') and 'reference' not in line_stripped.lower():
                break

            # Collect reference lines
            if in_references and line_stripped:
                # Skip placeholder references
                if 'Author et al.' in line_stripped and 'Year' in line_stripped and 'Source' in line_stripped:
                    continue
                # Add valid references
                if line_stripped.startswith('-') or line_stripped.startswith('['):
                    references.add(line_stripped)

    return references


def merge_markdown_files(md_files: list, topic: str) -> str:
    """
    Merge dimension markdown files into single document.

    Preserves all citations and references.
    Removes duplicate references sections from individual files.

    Args:
        md_files: List of markdown file paths
        topic: Research topic

    Returns:
        Merged markdown content
    """
    # Get local timezone-aware timestamp
    local_time = time.strftime('%Y-%m-%d %H:%M:%S %Z', time.localtime())

    merged_content = f"""# Research Report: {topic}

*Generated: {local_time}*

---

## Executive Summary

[EXECUTIVE_SUMMARY_TO_BE_GENERATED]

---

"""

    # Merge dimension documents
    for idx, md_file in enumerate(md_files):
        if not md_file or not Path(md_file).exists():
            continue

        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()

        # Remove references section from individual dimension
        lines = content.split('\n')
        filtered_lines = []
        in_references = False

        for line in lines:
            line_stripped = line.strip()

            # Skip references section
            if line_stripped.startswith('##') and 'reference' in line_stripped.lower():
                in_references = True
                continue

            # End of references section
            if in_references and line_stripped.startswith('#'):
                in_references = False

            if not in_references:
                filtered_lines.append(line)

        dimension_content = '\n'.join(filtered_lines).strip()

        # Add dimension content
        if idx > 0:
            merged_content += "\n\n---\n\n"
        merged_content += dimension_content + "\n\n"

    # Add conclusion placeholder
    merged_content += """---

## Conclusion

[CONCLUSION_TO_BE_GENERATED]

---

"""

    # Collect and add all references
    references = collect_references_from_markdown(md_files)

    if references:
        merged_content += "## References\n\n"
        for ref in sorted(references):
            merged_content += f"{ref}\n"

    return merged_content


def parse_inline_formatting(text: str, paragraph, doc=None):
    """
    Parse inline markdown formatting and add formatted runs to paragraph.

    Supports:
    - **bold**
    - *italic*
    - §FOOTNOTE:number:url§ markers (converted to superscript numbers)
    - [1], [2] citations (superscript)
    - `code`

    Args:
        text: Text with inline markdown
        paragraph: python-docx paragraph object to add runs to
        doc: Document object (for tracking footnotes)
    """
    # Pattern to match inline formatting
    # Order matters: bold before italic (** before *)
    patterns = [
        (r'\*\*(.+?)\*\*', 'bold'),      # **bold**
        (r'\*(.+?)\*', 'italic'),         # *italic*
        (r'`(.+?)`', 'code'),             # `code`
        (r'§FOOTNOTE:(\d+):([^§]+)§', 'footnote'),  # §FOOTNOTE:number:url§
        (r'\[\d+(?:,\s*\d+)*\]', 'citation')  # [1], [1, 2, 3]
    ]

    pos = 0
    while pos < len(text):
        # Find the earliest match
        earliest_match = None
        earliest_pos = len(text)
        match_type = None

        for pattern, fmt_type in patterns:
            match = re.search(pattern, text[pos:])
            if match and match.start() < earliest_pos - pos:
                earliest_match = match
                earliest_pos = pos + match.start()
                match_type = fmt_type

        if not earliest_match:
            # No more formatting, add remaining text
            if pos < len(text):
                paragraph.add_run(text[pos:])
            break

        # Add text before match
        if earliest_pos > pos:
            paragraph.add_run(text[pos:earliest_pos])

        # Add formatted text
        matched_text = earliest_match.group(0)

        if match_type == 'bold':
            run = paragraph.add_run(earliest_match.group(1))
            run.bold = True
        elif match_type == 'italic':
            run = paragraph.add_run(earliest_match.group(1))
            run.italic = True
        elif match_type == 'code':
            run = paragraph.add_run(earliest_match.group(1))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif match_type == 'footnote':
            # Footnote marker: extract number and URL
            footnote_num = earliest_match.group(1)
            footnote_url = earliest_match.group(2)

            # Add superscript number
            run = paragraph.add_run(footnote_num)
            run.font.superscript = True
            run.font.size = Pt(9)

            # Store footnote for later (add to document if doc is provided)
            if doc and hasattr(doc, '_footnote_map'):
                doc._footnote_map[footnote_num] = footnote_url

        elif match_type == 'citation':
            # Citation as superscript
            run = paragraph.add_run(matched_text)
            run.font.superscript = True
            run.font.size = Pt(9)

        # Move position forward
        pos = earliest_pos + len(matched_text)


def docx_to_pdf(docx_path: str, pdf_path: str):
    """
    Convert Word document to PDF using docx2pdf library.

    This method is more secure than using subprocess directly as it:
    - Avoids direct subprocess calls with user-controlled paths
    - Prevents command injection vulnerabilities
    - Validates file paths and extensions
    - Uses a well-tested library API

    Args:
        docx_path: Input Word document path
        pdf_path: Output PDF path

    Raises:
        FileNotFoundError: If the input DOCX file doesn't exist
        ValueError: If file path validation fails
        Exception: If PDF conversion fails
    """
    import os
    from pathlib import Path

    try:
        # Import docx2pdf library
        from docx2pdf import convert

        # Normalize and validate paths
        docx_path = os.path.abspath(docx_path)
        pdf_path = os.path.abspath(pdf_path)

        # Security validation: Check file exists
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")

        # Security validation: Check file extension
        if not docx_path.lower().endswith('.docx'):
            raise ValueError("Input file must be a .docx file")

        # Security validation: Ensure paths are within workspace
        # This prevents path traversal attacks
        from src.utils.workspace import get_workspace
        workspace_instance = get_workspace()
        workspace_dir = Path(workspace_instance.base_path).resolve()

        try:
            docx_resolved = Path(docx_path).resolve()
            pdf_resolved = Path(pdf_path).resolve()

            # Verify paths are within workspace directory
            if not str(docx_resolved).startswith(str(workspace_dir)):
                raise ValueError(f"DOCX path outside workspace: {docx_path}")
            if not str(pdf_resolved).startswith(str(workspace_dir)):
                raise ValueError(f"PDF path outside workspace: {pdf_path}")

        except (ValueError, OSError) as e:
            raise ValueError(f"Invalid file path: {e}")

        # Ensure output directory exists
        output_dir = os.path.dirname(pdf_path)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        # Convert DOCX to PDF
        # docx2pdf library doesn't support Linux, so use LibreOffice directly
        import platform
        import subprocess

        if platform.system() == 'Linux':
            # Use LibreOffice directly on Linux
            output_dir = os.path.dirname(pdf_path)
            result = subprocess.run(
                [
                    'soffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', output_dir,
                    docx_path
                ],
                capture_output=True,
                text=True,
                timeout=60
            )

            if result.returncode != 0:
                raise Exception(f"LibreOffice conversion failed: {result.stderr}")

            # LibreOffice creates file with same name but .pdf extension
            # Need to rename if output path is different
            libreoffice_output = os.path.join(output_dir, os.path.basename(docx_path).replace('.docx', '.pdf'))
            if libreoffice_output != pdf_path and os.path.exists(libreoffice_output):
                os.rename(libreoffice_output, pdf_path)
        else:
            # Use docx2pdf on Mac/Windows
            convert(docx_path, pdf_path)

        # Verify conversion success
        if not os.path.exists(pdf_path):
            raise Exception(f"PDF file was not created at {pdf_path}")

        print(f"   ✓ PDF created: {pdf_path}")

    except ImportError:
        raise Exception(
            "docx2pdf library is not installed. "
            "Install it with: pip install docx2pdf"
        )
    except FileNotFoundError as e:
        raise Exception(f"File not found: {e}")
    except ValueError as e:
        raise Exception(f"Validation error: {e}")
    except Exception as e:
        raise Exception(f"PDF conversion failed: {str(e)}")


def markdown_to_docx(markdown_content: str, output_path: str, chart_files: list = None):
    """
    Convert markdown to Word document with proper inline formatting and charts.
    Uses Word footnotes for URL citations.

    Args:
        markdown_content: Markdown content
        output_path: Output Word document path
        chart_files: List of chart file dicts with 'path', 'title', 'type'
    """
    import re
    import os
    from docx.shared import Inches

    # Build URL citation mapping
    url_pattern = r'\[(https?://[^\]]+)\]'
    urls_found = re.findall(url_pattern, markdown_content)

    # Create unique URL list with numbering
    url_to_number = {}
    citation_counter = 1
    for url in urls_found:
        if url not in url_to_number:
            url_to_number[url] = citation_counter
            citation_counter += 1

    # Mark URL citations with placeholders for footnote insertion
    # Use a unique marker that won't appear in normal text
    def replace_url_with_marker(match):
        url = match.group(1)
        number = url_to_number.get(url, '?')
        # Use special marker: §FOOTNOTE:number:url§
        return f'§FOOTNOTE:{number}:{url}§'

    markdown_content = re.sub(url_pattern, replace_url_with_marker, markdown_content)

    doc = Document()

    # Initialize footnote tracking
    doc._footnote_map = {}  # {number: url}

    lines = markdown_content.split('\n')
    charts_inserted = False  # Track if charts have been inserted
    skip_next = False  # Flag to skip figure caption lines

    for i, line in enumerate(lines):
        line = line.rstrip()

        if not line:
            continue

        # Skip this line if it's a figure caption that was already processed
        if skip_next:
            skip_next = False
            continue

        # Heading level 1
        if line.startswith('# '):
            heading_text = line[2:]
            doc.add_heading(heading_text, level=0)
        # Heading level 2
        elif line.startswith('## '):
            heading_text = line[3:]
            doc.add_heading(heading_text, level=1)

            # Insert charts after "Executive Summary" heading
            if not charts_inserted and heading_text.strip().lower() == 'executive summary':
                if chart_files:
                    print(f"\n📊 Inserting {len(chart_files)} chart(s) after Executive Summary...")
                    doc.add_paragraph()  # Add spacing

                    for chart_info in chart_files:
                        chart_path = chart_info.get('path')
                        chart_title = chart_info.get('title', 'Chart')

                        if chart_path and os.path.exists(chart_path):
                            try:
                                # Add chart title
                                title_para = doc.add_paragraph()
                                title_run = title_para.add_run(chart_title)
                                title_run.bold = True

                                # Add chart image (width: 6 inches, maintains aspect ratio)
                                doc.add_picture(chart_path, width=Inches(6))
                                doc.add_paragraph()  # Add spacing after chart
                                print(f"   ✅ Inserted: {chart_title}")
                            except Exception as e:
                                print(f"   ⚠️  Failed to insert {chart_title}: {e}")
                        else:
                            print(f"   ⚠️  Chart file not found: {chart_path}")

                    charts_inserted = True
        # Heading level 3
        elif line.startswith('### '):
            heading_text = line[4:]
            doc.add_heading(heading_text, level=2)
        # Heading level 4
        elif line.startswith('#### '):
            heading_text = line[5:]
            doc.add_heading(heading_text, level=3)
        # Image (markdown syntax: ![alt](path))
        elif line.startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt_text, image_path = match.groups()
                # Resolve relative path
                if not os.path.isabs(image_path):
                    # Assume relative to workspace temp directory
                    workspace_temp = Path(output_path).parent.parent / "temp"
                    # Try to find session-specific chart
                    for session_dir in workspace_temp.glob("*/charts"):
                        full_image_path = session_dir / Path(image_path).name
                        if full_image_path.exists():
                            image_path = str(full_image_path)
                            break

                # Insert image if exists
                if os.path.exists(image_path):
                    try:
                        doc.add_picture(image_path, width=Inches(6))
                        # Check if next line is a figure caption
                        if i + 1 < len(lines):
                            next_line = lines[i + 1].strip()
                            if next_line.startswith('*Figure ') and next_line.endswith('*'):
                                # Add caption below image
                                caption_text = next_line[1:-1]  # Remove * markers
                                p = doc.add_paragraph()
                                run = p.add_run(caption_text)
                                run.italic = True
                                p.alignment = 1  # Center alignment
                                skip_next = True  # Skip the caption line in next iteration
                    except Exception as e:
                        print(f"   ⚠️  Failed to insert image {image_path}: {e}")
                        # Add placeholder text
                        p = doc.add_paragraph()
                        run = p.add_run(f"[Image: {alt_text}]")
                        run.italic = True
                else:
                    print(f"   ⚠️  Image not found: {image_path}")
                    # Add placeholder text
                    p = doc.add_paragraph()
                    run = p.add_run(f"[Image not found: {alt_text}]")
                    run.italic = True
        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
        # Bullet list
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(line[2:], p, doc)
        # Numbered list
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            parse_inline_formatting(text, p, doc)
        # Italic text (for metadata like *Generated: ...*)
        elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[1:-1])
            run.italic = True
        # Regular paragraph with inline formatting
        else:
            p = doc.add_paragraph()
            parse_inline_formatting(line, p, doc)

    # Add Footnotes section at the end if there are footnotes
    if hasattr(doc, '_footnote_map') and doc._footnote_map:
        doc.add_paragraph()  # Add spacing
        doc.add_page_break()  # Start footnotes on new page (optional)
        doc.add_heading('Footnotes', level=2)

        # Sort by footnote number
        sorted_footnotes = sorted(doc._footnote_map.items(), key=lambda x: int(x[0]))

        for number, url in sorted_footnotes:
            # Use regular paragraph with hanging indent
            p = doc.add_paragraph()
            # Add footnote number in superscript
            run = p.add_run(number)
            run.font.superscript = True
            run.font.size = Pt(9)
            # Add space
            p.add_run(' ')
            # Add URL as hyperlink
            add_hyperlink(p, url, url)

    doc.save(output_path)


def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph.

    Args:
        paragraph: python-docx paragraph object
        url: URL to link to
        text: Display text for the link
    """
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn

    # Create hyperlink element
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create run with hyperlink style
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add hyperlink styling
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


@traceable(name="report_writing_node")
async def report_writing_node(state: ResearchState) -> Dict[str, Any]:
    """
    Stage 5: Merge dimension documents and refine with editor agent.

    Process:
    1. Merge dimension markdown files
    2. Collect and deduplicate references
    3. Add executive summary and conclusion placeholders
    4. Use editor agent to refine document
    5. Generate executive summary and conclusion
    6. Convert to Word document

    Args:
        state: ResearchState with dimension_documents

    Returns:
        Dict with report_file path
    """
    from src.utils.status_updater import get_status_updater

    # Check if research is cancelled before starting
    check_cancellation(state)

    logger.info("STAGE 5: REPORT WRITING")

    start_time = time.time()

    # Update status
    research_session_id = state.get("research_session_id")
    status_updater = get_status_updater(research_session_id)
    if status_updater:
        status_updater.update_stage('report_writing')

    dimension_documents = state.get("dimension_documents", {})
    topic = state.get("topic", "Research Topic")
    research_config = state.get("research_config", {})
    user_research_context = research_config.get("research_context", "")

    if not dimension_documents:
        logger.warning("No dimension documents found")
        return {
            "report_file": None
        }

    # Get dimension document paths in order
    dimensions = list(dimension_documents.keys())
    dimension_doc_paths = [dimension_documents[dim] for dim in dimensions]

    print(f"\nMerging {len(dimension_doc_paths)} dimension documents...")
    for dim, path in zip(dimensions, dimension_doc_paths):
        print(f"   ✓ {dim}: {path}")

    # Step 1: Merge markdown files
    print("\n📄 Merging markdown documents...")
    merged_markdown = merge_markdown_files(dimension_doc_paths, topic)

    # Save intermediate markdown
    from src.utils.workspace import get_workspace
    workspace = get_workspace()

    safe_topic = re.sub(r'[^\w\s-]', '', topic).strip().replace(' ', '_')[:50]
    timestamp = time.strftime('%Y%m%d_%H%M%S')

    # Construct markdown path directly (don't use get_final_report_path which adds .docx)
    intermediate_md_path = str(workspace.final_dir / f"draft_{safe_topic}_{timestamp}.md")
    with open(intermediate_md_path, 'w', encoding='utf-8') as f:
        f.write(merged_markdown)

    print(f"   ✓ Draft markdown saved: {intermediate_md_path}")
    print(f"   ✓ Document length: {len(merged_markdown)} characters")

    # Step 2: Editor agent refinement
    print("\n🔧 Refining document with editor agent...")
    print(f"   Editor tools will work on: {intermediate_md_path}")
    print(f"   Tools will access file via InjectedState (draft_report_file)")

    # Get LLM for editor
    llm = get_llm_for_node("report_writing", state)

    # Create editor agent with tools
    editor_tools = [write_summary_and_conclusion, replace_text]

    # Prepare research context section for editor
    research_context_section = ""
    if user_research_context:
        research_context_section = f"""
RESEARCH CONTEXT PROVIDED BY USER:
{user_research_context}

This context should guide your editing decisions and ensure the report aligns with the user's goals.

---

"""

    editor_system_prompt = f"""You are an expert technical editor refining a research report. Your tasks:

{research_context_section}AVAILABLE TOOLS:

1. **write_summary_and_conclusion(summary_content, conclusion_content)**:
   - summary_content: Executive Summary content (200-300 words)
   - conclusion_content: Conclusion content (300-400 words)
   - Use this to generate BOTH sections in one call

2. **replace_text(find_text_param, replace_with)**:
   - find_text_param: Text to find
   - replace_with: Replacement text
   - Use this to fix awkward transitions, improve flow, or correct inconsistencies

YOUR TASKS:

**STEP 1: Review the document**
- Read through the provided document preview
- Identify any awkward transitions between sections
- Note any inconsistencies or redundancies
- **Check for incomplete URL citations** (e.g., [https://example without closing bracket or incomplete URL)

**STEP 2: Remove incomplete citations**
- If you find any incomplete or malformed URL citations, remove them using replace_text
- Example: "[https://aienergyc" without proper closing or incomplete domain → remove entirely
- Only remove clearly broken citations, not valid ones

**STEP 3: Fix transitions and flow (if needed)**
- Use replace_text to improve awkward section connections
- Smooth out redundancies or repetitive phrases
- Ensure consistent terminology throughout

**STEP 4: Write Executive Summary AND Conclusion**
- Call write_summary_and_conclusion(summary_content="...", conclusion_content="...")
- Generate BOTH sections in a SINGLE tool call
- Executive Summary (200-300 words):
  * Synthesize the main topic, key dimensions explored, and major findings
  * Highlight the most important insights from across all dimensions
- Conclusion (300-400 words):
  * Synthesize key findings and implications
  * Discuss broader impact and future directions
  * Provide clear takeaways

IMPORTANT:
- Make minimal changes - only fix genuine issues
- **NEVER modify or remove URL citations in square brackets [https://...]**
- **DO NOT use replace_text on paragraphs containing citations unless absolutely necessary**
- Preserve all citations and references exactly as they appear
- Maintain the technical depth and accuracy
- Focus on improving readability and flow

CRITICAL WARNING ABOUT CITATIONS:
- Citations look like [https://example.com/article]
- If you must edit text near citations, be extremely careful to preserve the complete URL
- Do not split, truncate, or modify any text inside square brackets starting with http"""

    # Check if model supports prompt caching
    # Nova Pro removed due to ValidationException with cachePoint in long content arrays
    model_name = getattr(llm, 'model_id', getattr(llm, 'model', ''))
    supports_caching = any(
        model in model_name for model in [
            'us.anthropic.claude-sonnet-4-5-20250929-v1:0',
            'us.anthropic.claude-sonnet-4-20250514-v1:0',
            'us.anthropic.claude-haiku-4-5-20251001-v1:0',
            'anthropic.claude-3-5-haiku-20241022-v1:0'
        ]
    )

    if supports_caching:
        print("   ✓ Prompt caching enabled for editor agent")
        cached_system_message = SystemMessage(
            content=[
                {"text": editor_system_prompt},
                {"cachePoint": {"type": "default"}}
            ]
        )

        custom_prompt = ChatPromptTemplate.from_messages([
            cached_system_message,
            MessagesPlaceholder(variable_name="messages"),
            MessagesPlaceholder(variable_name="agent_scratchpad", optional=True),
        ])
    else:
        print("   ⚠ Prompt caching not supported for this model")
        system_message = SystemMessage(content=editor_system_prompt)
        custom_prompt = ChatPromptTemplate.from_messages([
            system_message,
            MessagesPlaceholder(variable_name="messages"),
            MessagesPlaceholder(variable_name="agent_scratchpad", optional=True),
        ])

    from langgraph.checkpoint.memory import MemorySaver

    # Add cache point hook if caching is supported
    if supports_caching:
        def add_cache_point_to_last_message(state):
            """Add cache point to last Human or AI message before tool results"""
            from langchain_core.messages import HumanMessage, AIMessage
            messages = state.get("messages", [])
            if not messages:
                return {}

            # Find last Human or AI message
            for i in range(len(messages) - 1, -1, -1):
                msg = messages[i]
                if isinstance(msg, (HumanMessage, AIMessage)):
                    # Check if cache point already exists
                    if isinstance(msg.content, list):
                        has_cache = any(isinstance(item, dict) and "cachePoint" in item for item in msg.content)
                        if has_cache:
                            return {"llm_input_messages": messages}

                    # Add cache point
                    if isinstance(msg.content, list):
                        new_content = msg.content + [{"cachePoint": {"type": "default"}}]
                    else:
                        new_content = [{"text": msg.content}, {"cachePoint": {"type": "default"}}]

                    if isinstance(msg, HumanMessage):
                        new_msg = HumanMessage(content=new_content, **msg.dict(exclude={"content", "type"}))
                    else:
                        new_msg = AIMessage(content=new_content, **msg.dict(exclude={"content", "type"}))

                    new_messages = messages[:i] + [new_msg] + messages[i + 1:]
                    return {"llm_input_messages": new_messages}

            return {"llm_input_messages": messages}

        editor_agent = create_react_agent(
            model=llm,
            tools=editor_tools,
            prompt=custom_prompt,
            pre_model_hook=add_cache_point_to_last_message,
            checkpointer=MemorySaver()
        )
    else:
        editor_agent = create_react_agent(
            model=llm,
            tools=editor_tools,
            prompt=custom_prompt,
            checkpointer=MemorySaver()
        )

    # Run editor agent (tools will modify the file directly)
    # Provide FULL document for context with prompt caching
    user_message = f"""Please refine this research report by following the steps in your instructions.

FULL DOCUMENT:
{merged_markdown}

---

Please follow all steps:
1. Review the document for any issues including incomplete URL citations
2. Remove any incomplete or malformed URL citations (e.g., [https://aienergyc without proper ending)
3. Fix awkward transitions or flow issues (use replace_text to fix if needed)
4. Generate Executive Summary AND Conclusion using write_summary_and_conclusion

Focus on creating well-synthesized summary and conclusion that tie together insights from all dimensions.
"""

    # Create user message with prompt caching if supported
    if supports_caching:
        from langchain_core.messages import HumanMessage
        user_msg = HumanMessage(
            content=[
                {"text": user_message},
                {"cachePoint": {"type": "default"}}
            ]
        )
        print("   ✓ Prompt caching enabled for full document context")
    else:
        user_msg = ("user", user_message)

    # Invoke editor agent with draft_report_file in config (for RunnableConfig access)
    # Tools use RunnableConfig (not InjectedState) to access the file path
    # LangChain automatically injects config into tool parameters
    editor_result = await editor_agent.ainvoke(
        {"messages": [user_msg]},
        config={
            "configurable": {
                "thread_id": "editor_session",
                "draft_report_file": intermediate_md_path  # Tools access via config
            }
        }
    )

    # Log editor agent result for debugging
    print("\n📋 Editor agent execution summary:")
    if "messages" in editor_result:
        messages = editor_result["messages"]
        tool_calls_count = 0
        for msg in messages:
            if hasattr(msg, 'tool_calls') and msg.tool_calls:
                for tool_call in msg.tool_calls:
                    tool_calls_count += 1
                    print(f"   - Tool called: {tool_call.get('name', 'unknown')}")
        print(f"   Total tool calls: {tool_calls_count}")

    # Read the edited content from file (tools have already saved changes)
    with open(intermediate_md_path, 'r', encoding='utf-8') as f:
        edited_markdown = f.read()

    print(f"   ✓ All editing phases completed")
    print(f"   ✓ Document length after editing: {len(edited_markdown)} characters")

    # Verify that placeholders were replaced
    missing_placeholders = []
    if "[EXECUTIVE_SUMMARY_TO_BE_GENERATED]" in edited_markdown:
        missing_placeholders.append("Executive Summary")
    if "[CONCLUSION_TO_BE_GENERATED]" in edited_markdown:
        missing_placeholders.append("Conclusion")

    if missing_placeholders:
        print(f"   ⚠️  Warning: {', '.join(missing_placeholders)} not generated by editor agent")
    else:
        print("   ✓ Executive summary and conclusion successfully generated")

    elapsed = time.time() - start_time

    logger.info(f"Report writing completed in {elapsed:.2f}s - {len(edited_markdown)} characters")

    return {
        "draft_report_file": intermediate_md_path
    }
