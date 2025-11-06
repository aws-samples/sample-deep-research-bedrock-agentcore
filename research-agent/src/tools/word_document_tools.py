"""Word Document Writing Tools

Tools for AI agents to create structured Word documents.
Each tool performs a specific formatting action on a Word document.
"""

from typing import Optional, List, Dict, Any
from pydantic import BaseModel, Field
from langchain_core.tools import BaseTool
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement


# Global document storage (keyed by document_id)
# Each parallel dimension_reducer uses unique document_id (timestamp + uuid)
# so no race conditions
_active_documents: Dict[str, Document] = {}


class CreateDocumentInput(BaseModel):
    """Input for creating a new document"""
    document_id: str = Field(description="Unique identifier for this document (e.g., 'dimension_1')")
    title: Optional[str] = Field(default=None, description="Optional document title")


class AddHeadingInput(BaseModel):
    """Input for adding a heading"""
    document_id: str = Field(description="Document identifier")
    text: str = Field(description="Heading text")
    level: int = Field(default=2, description="Heading level (1=largest, 2=section, 3=subsection)")


class AddParagraphInput(BaseModel):
    """Input for adding a paragraph"""
    document_id: str = Field(description="Document identifier")
    text: str = Field(description="Paragraph text content")
    style: Optional[str] = Field(default=None, description="Optional style name (e.g., 'List Bullet', 'Intense Quote')")


class AddBulletListInput(BaseModel):
    """Input for adding a bullet list"""
    document_id: str = Field(description="Document identifier")
    items: List[str] = Field(description="List of bullet point items")


class AddTableInput(BaseModel):
    """Input for adding a table"""
    document_id: str = Field(description="Document identifier")
    headers: List[str] = Field(description="Column headers")
    rows: List[List[str]] = Field(description="Table rows (each row is a list of cell values)")


class AddCitationInput(BaseModel):
    """Input for adding a citation"""
    document_id: str = Field(description="Document identifier")
    citation_text: str = Field(description="Citation text (e.g., 'Smith et al., 2024, arXiv:2401.12345')")
    context: Optional[str] = Field(default=None, description="Optional context text before citation")


class AddPageBreakInput(BaseModel):
    """Input for adding a page break"""
    document_id: str = Field(description="Document identifier")


class SaveDocumentInput(BaseModel):
    """Input for saving a document"""
    document_id: str = Field(description="Document identifier")
    filename: str = Field(description="Output filename (e.g., 'dimension_1.docx')")


class CreateDocumentTool(BaseTool):
    """Create a new Word document"""
    name: str = "create_document"
    description: str = """Create a new Word document to write content into.

    Returns the document_id that should be used in all subsequent tool calls.

    Example: create_document(document_id="dimension_interoperability", title="Interoperability and Integration")
    """
    args_schema: type[BaseModel] = CreateDocumentInput

    def _run(self, document_id: str, title: Optional[str] = None) -> str:
        """Create new document"""
        # Check if document already exists
        if document_id in _active_documents:
            return f"⚠️ Warning: Document '{document_id}' already exists. Overwriting."

        doc = Document()

        if title:
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        _active_documents[document_id] = doc

        return f"✅ Document '{document_id}' created{' with title: ' + title if title else ''}"


class AddHeadingTool(BaseTool):
    """Add a heading to the document"""
    name: str = "add_heading"
    description: str = """Add a heading to the document.

    Use level 2 (##) for main section headings.
    Use level 3 (###) for subsections.

    Example: add_heading(document_id="dimension_1", text="Standardization Challenges", level=2)
    """
    args_schema: type[BaseModel] = AddHeadingInput

    def _run(self, document_id: str, text: str, level: int = 2) -> str:
        """Add heading"""
        if document_id not in _active_documents:
            return f"❌ Error: Document '{document_id}' not found. Create it first with create_document."

        doc = _active_documents[document_id]
        doc.add_heading(text, level=level)

        return f"✅ Added heading: {text}"


class AddParagraphTool(BaseTool):
    """Add a paragraph to the document"""
    name: str = "add_paragraph"
    description: str = """Add a paragraph of text to the document.

    The text can include citations in format [Author et al., Year, Source].
    These will be automatically formatted in blue italic.

    Example: add_paragraph(document_id="dimension_1", text="Recent research shows... [Smith et al., 2024, arXiv:2401.12345]")
    """
    args_schema: type[BaseModel] = AddParagraphInput

    def _run(self, document_id: str, text: str, style: Optional[str] = None) -> str:
        """Add paragraph"""
        if document_id not in _active_documents:
            return f"❌ Error: Document '{document_id}' not found"

        doc = _active_documents[document_id]
        para = doc.add_paragraph()

        if style:
            para.style = style

        # Parse citations and format them
        import re
        pattern = r'\[([^\]]+(?:et al\.|[A-Z][a-z]+)[^\]]*(?:19|20)\d{2}[^\]]*)\]'

        parts = []
        last_end = 0

        for match in re.finditer(pattern, text):
            # Add text before citation
            if match.start() > last_end:
                parts.append(('normal', text[last_end:match.start()]))

            # Add citation
            parts.append(('citation', match.group(0)))
            last_end = match.end()

        # Add remaining text
        if last_end < len(text):
            parts.append(('normal', text[last_end:]))

        # Add runs with appropriate formatting
        for format_type, content in parts:
            run = para.add_run(content)
            if format_type == 'citation':
                run.italic = True
                run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

        return f"✅ Added paragraph ({len(text)} chars)"


class AddBulletListTool(BaseTool):
    """Add a bullet list to the document"""
    name: str = "add_bullet_list"
    description: str = """Add a bullet list to the document.

    Example: add_bullet_list(document_id="dimension_1", items=["First point", "Second point", "Third point"])
    """
    args_schema: type[BaseModel] = AddBulletListInput

    def _run(self, document_id: str, items: List[str]) -> str:
        """Add bullet list"""
        if document_id not in _active_documents:
            return f"❌ Error: Document '{document_id}' not found"

        doc = _active_documents[document_id]

        for item in items:
            doc.add_paragraph(item, style='List Bullet')

        return f"✅ Added bullet list with {len(items)} items"


class AddTableTool(BaseTool):
    """Add a table to the document"""
    name: str = "add_table"
    description: str = """Add a table to the document.

    Example:
    add_table(
        document_id="dimension_1",
        headers=["Method", "Accuracy", "Speed"],
        rows=[
            ["Method A", "95%", "Fast"],
            ["Method B", "97%", "Slow"]
        ]
    )
    """
    args_schema: type[BaseModel] = AddTableInput

    def _run(self, document_id: str, headers: List[str], rows: List[List[str]]) -> str:
        """Add table"""
        if document_id not in _active_documents:
            return f"❌ Error: Document '{document_id}' not found"

        doc = _active_documents[document_id]

        # Create table
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Light Grid Accent 1'

        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add data rows
        for i, row_data in enumerate(rows, 1):
            cells = table.rows[i].cells
            for j, cell_value in enumerate(row_data):
                cells[j].text = str(cell_value)

        return f"✅ Added table ({len(headers)} columns × {len(rows)} rows)"


class AddCitationTool(BaseTool):
    """Add an inline citation with special formatting"""
    name: str = "add_citation"
    description: str = """Add an inline citation in the current paragraph or as a new citation paragraph.

    Citations are formatted in blue italic for easy identification.

    Example: add_citation(document_id="dimension_1", citation_text="Smith et al., 2024, arXiv:2401.12345", context="Recent research shows that")

    This creates: "Recent research shows that [Smith et al., 2024, arXiv:2401.12345]"
    """
    args_schema: type[BaseModel] = AddCitationInput

    def _run(self, document_id: str, citation_text: str, context: Optional[str] = None) -> str:
        """Add citation"""
        if document_id not in _active_documents:
            return f"❌ Error: Document '{document_id}' not found"

        doc = _active_documents[document_id]

        # Create Citation style if it doesn't exist
        try:
            citation_style = doc.styles['Citation']
        except KeyError:
            from docx.enum.style import WD_STYLE_TYPE
            citation_style = doc.styles.add_style('Citation', WD_STYLE_TYPE.PARAGRAPH)
            citation_style.base_style = doc.styles['Normal']
            font = citation_style.font
            font.italic = True
            font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

        # Add paragraph with citation
        para = doc.add_paragraph(style='Citation')

        if context:
            # Add context as normal text
            normal_run = para.add_run(context + " ")
            normal_run.italic = False
            normal_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        # Add citation in brackets with blue italic
        citation_run = para.add_run(f"[{citation_text}]")
        citation_run.italic = True
        citation_run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

        return f"✅ Added citation: [{citation_text}]"


class AddPageBreakTool(BaseTool):
    """Add a page break"""
    name: str = "add_page_break"
    description: str = """Add a page break to start content on a new page.

    Use this sparingly, typically only between major sections.

    Example: add_page_break(document_id="dimension_1")
    """
    args_schema: type[BaseModel] = AddPageBreakInput

    def _run(self, document_id: str) -> str:
        """Add page break"""
        if document_id not in _active_documents:
            return f"❌ Error: Document '{document_id}' not found"

        doc = _active_documents[document_id]
        doc.add_page_break()

        return f"✅ Added page break"


class SaveDocumentTool(BaseTool):
    """Save the document to a file"""
    name: str = "save_document"
    description: str = """Save the document to a Word file.

    This should be called once at the end when the document is complete.

    Example: save_document(document_id="dimension_1", filename="dimension_interoperability.docx")
    """
    args_schema: type[BaseModel] = SaveDocumentInput

    def _run(self, document_id: str, filename: str) -> str:
        """Save document"""
        if document_id not in _active_documents:
            return f"❌ Error: Document '{document_id}' not found"

        doc = _active_documents[document_id]
        doc.save(filename)

        return f"✅ Document saved to: {filename}"


# Create tool instances
create_document_tool = CreateDocumentTool()
add_heading_tool = AddHeadingTool()
add_paragraph_tool = AddParagraphTool()
add_bullet_list_tool = AddBulletListTool()
add_table_tool = AddTableTool()
add_citation_tool = AddCitationTool()
add_page_break_tool = AddPageBreakTool()
save_document_tool = SaveDocumentTool()

# Tool list for agents
word_document_tools = [
    create_document_tool,
    add_heading_tool,
    add_paragraph_tool,
    add_bullet_list_tool,
    add_table_tool,
    add_citation_tool,
    add_page_break_tool,
    save_document_tool,
]


def clear_active_documents():
    """Clear all active documents (useful for testing)"""
    global _active_documents
    _active_documents = {}
